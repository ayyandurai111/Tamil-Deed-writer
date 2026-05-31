"""
tools/generate_docx.py
======================
Professional Tamil Sale Deed layout:
- Single uniform font: Latha 12pt throughout
- Title only: 14pt Bold Center
- "சொத்து விவரம்" label: 12pt Bold
- No section headings, no dividers, no bold sub-headers in body
- Date + Purchaser + Vendor in ONE opening paragraph
- All clauses as justified paragraphs
- Property as one comma-separated paragraph
- Line-by-line only for witnesses and signatures

CLEANUP RULES (optional fields):
- Any value that is None, "", or starts with "{{" → skipped entirely
- Label is NOT printed when value is absent
- Grammar fix already applied by fill_skeleton Phase 3
"""

import json
from datetime import datetime
from pathlib import Path
from mcp.types import Tool, TextContent

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from constants import OUTPUT_DIR
from file_store import put as _mem_store

# ── Tool definition ────────────────────────────────────────────────────────────
TOOL_DEFINITION = Tool(
    name="generate_docx",
    description=(
        "[CALL 7 of 8] ONE TASK: call this tool only. "
        "Pass: filled_skeleton = CALL 6 'clean_skeleton'. "
        "filename_prefix = 'vendorname_purchasername' format (e.g. 'ramasamy_murugan'). "
        "After tool returns: "
        "success=True  → show PAN/TDS notes if any, show legal disclaimer. Next: CALL 8 list_output_files. "
        "success=False → '❌ தோல்வி: [error] — மீண்டும் முயற்சிக்கவும்.'"
    ),
    inputSchema={
        "type": "object",
        "properties": {
            "filled_skeleton": {
                "type": "object",
                "description": "The filled skeleton JSON returned by fill_skeleton."
            },
            "filename_prefix": {
                "type": "string",
                "description": "Optional prefix for the output filename (e.g. 'raman_karthik').",
                "default": "deed"
            }
        },
        "required": ["filled_skeleton"],
        "additionalProperties": False
    },
    outputSchema={
        "type": "object",
        "properties": {
            "success":   {"type": "boolean"},
            "filename":  {"type": ["string", "null"]},
            "file":      {"type": ["string", "null"]},
            "error":     {"type": ["string", "null"]},
            "message":   {"type": "string"},
            "next_tool": {"type": ["string", "null"]}
        },
        "required": ["success", "filename", "file", "error", "message", "next_tool"],
        "additionalProperties": False
    },
    annotations={
        "title":           "DOCX Generator",
        "readOnlyHint":    False,
        "destructiveHint": False,
        "idempotentHint":  False,
    }
)


# ══════════════════════════════════════════════════════════════════════════════
#  CORE FONT HELPER
# ══════════════════════════════════════════════════════════════════════════════

BODY_FONT  = "Latha"
BODY_SIZE  = 12
TITLE_SIZE = 14


def _apply_font(run, size_pt: int = BODY_SIZE, bold: bool = False):
    run.font.name = "Latha"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BODY_FONT)
    rPr.insert(0, rFonts)


def _val(v, fallback: str = "") -> str:
    """
    Return clean string value, or fallback if absent/unfilled.
    Default fallback is "" — callers decide whether to skip or use a placeholder.
    Never returns "___________" automatically.
    """
    if v is None:
        return fallback
    s = str(v).strip()
    if s == "" or s.startswith("{{"):
        return fallback
    return s


def _req(v) -> str:
    """Required field — returns value or '___________' placeholder."""
    return _val(v, "___________")


# ══════════════════════════════════════════════════════════════════════════════
#  PARAGRAPH BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def _title_para(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    _apply_font(run, size_pt=TITLE_SIZE, bold=True)
    return p


def _body_para(doc, text: str, bold: bool = False,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY,
               space_before: int = 6, space_after: int = 6,
               first_line_indent: bool = True):
    if not text or str(text).startswith("{{"):
        return None
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    _apply_font(run, size_pt=BODY_SIZE, bold=bold)
    return p


def _sig_para(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    _apply_font(run, size_pt=BODY_SIZE, bold=False)
    return p


def _spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run("")
    _apply_font(run)
    return p


# ══════════════════════════════════════════════════════════════════════════════
#  PARTY TEXT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def _party_text(d: dict, deed_type: str = "plot") -> str:
    """
    Build full party description as flowing Tamil prose.
    Optional fields (caste, occupation, pan) are included only if present.
    """
    name     = _req(d.get("name"))
    age      = _req(d.get("age"))
    father   = _req(d.get("father_name"))
    relation = _val(d.get("relation"), "மகன்")
    address  = _req(d.get("address"))
    village  = _val(d.get("village"))
    vattam   = _val(d.get("vattam"))
    district = _val(d.get("district"))
    aadhaar  = _req(d.get("aadhaar", d.get("id_card", "")))
    phone    = _req(d.get("phone"))
    prefix   = _val(d.get("prefix"), "திரு").rstrip(".")

    # Build address — avoid duplication
    addr_parts = [address]
    if village and village not in address:
        addr_parts.append(village)
    if vattam and vattam not in address:
        addr_parts.append(vattam)
    if district and district not in address:
        addr_parts.append(district)
    addr_str = ", ".join(p for p in addr_parts if p)

    # Build identity
    if deed_type == "agriculture":
        caste      = _val(d.get("caste"))
        occupation = _val(d.get("occupation"))
        id_part = f"{prefix}.{name}"
        if father and not father.startswith("_"):
            id_part += f" அவர்களின் தந்தை {father}"
        id_part += f" சுமார் {age} வயதுள்ள"
        if caste:
            id_part += f" {caste}"
        if occupation:
            id_part += f" {occupation}"
    else:
        id_part = f"{prefix}.{father} அவர்களின் {relation} சுமார் {age} வயதுள்ள {prefix}.{name}"

    pan = _val(d.get("pan"))

    text = f"{addr_str} விலாசத்தில் வசிக்கும் {id_part}"
    text += f" (அடையாள அட்டை {aadhaar})"
    if pan:
        text += f" (PAN: {pan})"
    text += f" (கைபேசி எண்.{phone})"
    return text


# ══════════════════════════════════════════════════════════════════════════════
#  AGRICULTURE DEED BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def _build_agriculture_docx(data: dict, output_path: Path):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.5)
        sec.right_margin  = Cm(2.5)

    hdr  = data.get("header",        {})
    v    = data.get("vendor",        {})
    pur  = data.get("purchaser",     {})
    prop = data.get("property",      {})
    con  = data.get("consideration", {})
    wit  = data.get("witnesses",     [{}, {}])

    # ── 1. STAMP NOTE ────────────────────────────────────────────────────────
    stamp = _val(data.get("stamp_note"))
    if stamp:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(stamp)
        _apply_font(run, size_pt=BODY_SIZE)

    # ── 2. TITLE ──────────────────────────────────────────────────────────────
    _title_para(doc, data.get("title", "சுத்த விக்கிரயப் பத்திரம்"))
    _title_para(doc, data.get("subtitle", "ABSOLUTE SALE DEED — AGRICULTURE LAND"))

    act_ref = _val(data.get("act_ref"))
    if act_ref:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(10)
        run = p.add_run(act_ref)
        _apply_font(run, size_pt=BODY_SIZE)

    # ── 3. OPENING PARAGRAPH ─────────────────────────────────────────────────
    date_day   = _req(hdr.get("date_day"))
    date_month = _req(hdr.get("date_month"))
    date_year  = _req(hdr.get("date_year"))
    reg_office = _val(hdr.get("registration_office"))
    taluk      = _val(hdr.get("taluk"))
    district   = _val(hdr.get("district"))

    date_str = f"{date_day}ம் தேதி {date_month} மாதம் {date_year}ம் ஆண்டு"
    if reg_office:
        date_str += f", {reg_office} சார்பதிவக அலுவலகம்"
    if taluk:
        date_str += f", {taluk} தாலுக்கா"
    if district:
        date_str += f", {district} மாவட்டம்"

    v_text   = _party_text(v,   deed_type="agriculture")
    pur_text = _party_text(pur, deed_type="agriculture")

    opening = (
        f"{date_str},\n\n"
        f"{v_text} ஆகிய நான்,\n\n"
        f"{pur_text} அவர்களுக்கு அடியிற்கண்ட சாட்சிகள் முன்னிலையில் "
        f"மனப்பூர்வமாய் சம்மதித்து எழுதிக் கொடுத்த விவசாய நில "
        f"சுத்த விக்கிரயப் பத்திரம் என்னவென்றால்,"
    )
    _body_para(doc, opening, space_before=8, space_after=6, first_line_indent=False)

    # ── 4. VENDOR DECLARATION ────────────────────────────────────────────────
    _body_para(doc, _val(data.get("section_4_text")))

    # ── 5. CONSIDERATION ──────────────────────────────────────────────────────
    total    = _req(con.get("total_amount"))
    words    = _req(con.get("amount_in_words"))
    advance  = _val(con.get("advance_amount"))
    adv_dt   = _val(con.get("advance_date"))
    balance  = _val(con.get("balance_amount"))
    bal_dt   = _val(con.get("balance_date"))
    pay_mode = _val(con.get("payment_mode"))
    txn_no   = _val(con.get("transaction_no"))
    txn_dt   = _val(con.get("transaction_date"))
    bank     = _val(con.get("bank_name"))

    words_clean = words.rstrip().removesuffix("மட்டும்").rstrip()
    con_text = (
        f"இந்த விவசாய நிலத்தை விக்கிரயம் செய்வதற்கு நிர்ணயிக்கப்பட்ட "
        f"மொத்த விலை ரூபாய் {total} (எழுத்தால்: {words_clean} மட்டும்)."
    )
    if advance:
        con_text += f" இதில் முன்பணமாக ரூபாய் {advance}"
        if adv_dt:
            con_text += f" ({adv_dt} தேதி)"
        con_text += " பெறப்பட்டது."
    if balance:
        con_text += f" இருப்பு தொகை ரூபாய் {balance}"
        if bal_dt:
            con_text += f" ({bal_dt} தேதி)"
        con_text += " பெறப்பட்டது."
    if pay_mode:
        con_text += f" செலுத்திய முறை: {pay_mode}."
    if txn_no:
        con_text += f" Transaction எண். {txn_no}"
        if txn_dt:
            con_text += f" தேதி {txn_dt}"
        if bank:
            con_text += f", {bank} வங்கி"
        con_text += "."
    _body_para(doc, con_text)

    # ── 6. PROPERTY SCHEDULE ──────────────────────────────────────────────────
    _body_para(doc, "சொத்து விவரம்", bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=8, space_after=2, first_line_indent=False)

    prop_parts = []
    for label, key in [
        ("ஜில்லா",          "district"),
        ("தாலுக்கா",         "taluk"),
        ("கிராமம்",          "village"),
        ("வருவாய் கிராமம்",  "revenue_village"),
        ("வட்டம்",           "vattam"),
        ("சர்வே எண்",        "survey_no"),
        ("உட்பிரிவு",         "subdivision"),
        ("பட்டா எண்",         "patta_no"),
        ("சிட்டா எண்",        "chitta_no"),       # optional — skipped if None
        ("A-Register எண்",   "a_register_no"),    # optional — skipped if None
        ("நில வகை",          "land_type"),
        ("நஞ்சை/புஞ்சை",     "land_nature"),
        ("நீர் ஆதாரம்",       "water_source"),
    ]:
        val = _val(prop.get(key))
        if val:
            prop_parts.append(f"{label}: {val}")

    extent_acre = _val(prop.get("extent_acre"))
    extent_cent = _val(prop.get("extent_cent"), "0")
    if extent_acre:
        prop_parts.append(f"மொத்த பரப்பளவு: {extent_acre} ஏக்கர் {extent_cent} சென்ட்")

    for label, key in [
        ("கிழக்கு எல்லை", "boundary_east"),
        ("மேற்கு எல்லை",  "boundary_west"),
        ("வடக்கு எல்லை",  "boundary_north"),
        ("தெற்கு எல்லை",  "boundary_south"),
    ]:
        val = _val(prop.get(key))
        if val:
            prop_parts.append(f"{label}: {val}")

    # Optional property extras — label+value only printed when value exists
    for label, key in [
        ("கட்டடங்கள்",       "buildings"),        # optional
        ("மரங்கள்",           "trees"),            # optional
        ("நீர் கட்டமைப்பு",   "water_structures"), # optional
    ]:
        val = _val(prop.get(key))
        if val:
            prop_parts.append(f"{label}: {val}")

    if prop_parts:
        _body_para(doc, ", ".join(prop_parts) + ".",
                   space_before=2, space_after=6, first_line_indent=False)

    # ── 7. CHAIN OF TITLE ─────────────────────────────────────────────────────
    chain = data.get("chain_of_title", [])
    if chain and isinstance(chain, list):
        chain_parts = []
        for entry in chain:
            if isinstance(entry, dict):
                owner_raw = entry.get("owner", "")
                if not owner_raw or str(owner_raw).strip() == "" \
                        or str(owner_raw).startswith("{{"):
                    continue
                owner  = str(owner_raw).strip()
                label  = entry.get("label", "")
                doc_no = _val(entry.get("doc_no"))
                part   = f"{label}: {owner}"
                if doc_no:
                    part += f" (ஆவண எண்: {doc_no})"
                chain_parts.append(part)
        if chain_parts:
            _body_para(doc, "முந்தைய உரிமைத் தொடர் — " + "; ".join(chain_parts) + ".")

    # ── 8. LEGAL CLAUSES ──────────────────────────────────────────────────────
    legal = data.get("legal_clauses", {})
    if isinstance(legal, dict):
        for clause in legal.values():
            _body_para(doc, _val(clause))
    elif isinstance(legal, list):
        for clause in legal:
            _body_para(doc, str(clause))

    # ── 9. AGRICULTURE SPECIAL ────────────────────────────────────────────────
    agri = data.get("agriculture_special", {})
    if isinstance(agri, dict):
        for key in ["land_nature_clause", "irrigation_clause", "fmb_clause", "adangal_clause"]:
            _body_para(doc, _val(agri.get(key)))

        # Optional: standing crops / trees detail / farm structure
        # Each printed only if value exists — label skipped when absent
        detail_parts = []
        for label, key in [
            ("நிலத்தில் உள்ள பயிர்கள்", "standing_crops"),   # optional
            ("மரங்கள் விவரம்",           "trees_detail"),     # optional
            ("பண்ணை கட்டமைப்பு",         "farm_structure"),   # optional
        ]:
            val = _val(agri.get(key))
            if val:
                detail_parts.append(f"{label}: {val}")
        if detail_parts:
            _body_para(doc, ", ".join(detail_parts) + ".")

    elif isinstance(agri, list):
        for item in agri:
            _body_para(doc, str(item))

    # ── 10. POSSESSION & TRANSFER ─────────────────────────────────────────────
    _body_para(doc, _val(data.get("section_8_text")))

    # ── 11. DOCUMENTS HANDED ──────────────────────────────────────────────────
    docs = data.get("documents_handed", {})
    doc_labels = {
        "mother_deed":    "தாய் பத்திரம்",
        "patta_copy":     "பட்டா நகல்",
        "chitta_adangal": "சிட்டா / அடங்கல்",
        "ec_copy":        "வில்லங்கச் சான்று (EC)",
        "fmb_sketch":     "FMB Sketch",
        "tax_receipts":   "வரி ரசீது",
        "id_copies":      "அடையாள ஆவண நகல்கள்",
        "other_docs":     "இதர ஆவணங்கள்",   # optional — skipped if None
    }

    if isinstance(docs, dict):
        # Auto-derive mother_deed from chain_of_title if not supplied
        mother = _val(docs.get("mother_deed"))
        if not mother:
            chain_docs = []
            for entry in data.get("chain_of_title", []):
                if isinstance(entry, dict):
                    if entry.get("label", "").startswith("தற்போதைய"):
                        continue
                    dn = _val(entry.get("doc_no"))
                    if dn and "/" in dn:
                        chain_docs.append(dn)
            if chain_docs:
                docs = dict(docs)
                docs["mother_deed"] = " மற்றும் ".join(chain_docs)

        # Default boolean doc fields to "ஆம்" if blank
        for bool_key in ("patta_copy", "chitta_adangal", "ec_copy",
                         "fmb_sketch", "tax_receipts", "id_copies"):
            if not _val(docs.get(bool_key)):
                docs = dict(docs)
                docs[bool_key] = "ஆம்"

    handed_parts = []
    if isinstance(docs, dict):
        for key, label in doc_labels.items():
            val = _val(docs.get(key))
            if val:                          # None / "" / missing → skip label entirely
                handed_parts.append(f"{label}: {val}")

    if handed_parts:
        _body_para(doc, "ஒப்படைக்கப்பட்ட ஆவணங்கள் — " + ", ".join(handed_parts) + ".")
    else:
        _body_para(doc, "மேற்கண்ட சொத்திற்கான அனைத்து அசல் ஆவணங்களும் "
                        "கொள்முதலாளரிடம் ஒப்படைக்கப்பட்டன.")

    # ── 12. CLOSING TEXT ──────────────────────────────────────────────────────
    _body_para(doc, _val(data.get("closing_text")), space_before=8, space_after=8)

    # ── 13. WITNESSES ─────────────────────────────────────────────────────────
    w1 = wit[0] if len(wit) > 0 else {}
    w2 = wit[1] if len(wit) > 1 else {}

    for idx, w in enumerate([w1, w2], start=1):
        label   = w.get("label", f"சாட்சி {idx}")
        name    = _req(w.get("name"))
        address = _req(w.get("address"))
        aadhaar = _val(w.get("aadhaar"))          # optional — skipped if None
        line    = f"{label} : {name}, {address}"
        if aadhaar:
            line += f", ஆதார்: {aadhaar}"
        line += ",   கையொப்பம் : _______________"
        _sig_para(doc, line)

    _spacer(doc)

    # ── 14. SIGNATURES ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f"விற்பனையாளர் : {_req(v.get('name'))}")
    _apply_font(run, size_pt=BODY_SIZE)
    tab = p.add_run("\t\t\t\t\t")
    _apply_font(tab, size_pt=BODY_SIZE)
    run2 = p.add_run(f"கொள்முதலாளர் : {_req(pur.get('name'))}")
    _apply_font(run2, size_pt=BODY_SIZE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.first_line_indent = Cm(0)
    s1 = p2.add_run("கையொப்பம் / கட்டை விரல் ரேகை : _______________")
    _apply_font(s1, size_pt=BODY_SIZE)
    tab2 = p2.add_run("\t\t\t")
    _apply_font(tab2, size_pt=BODY_SIZE)
    s2 = p2.add_run("கையொப்பம் / கட்டை விரல் ரேகை : _______________")
    _apply_font(s2, size_pt=BODY_SIZE)

    # ── 15. REGISTRAR NOTE ────────────────────────────────────────────────────
    reg_note = _val(data.get("registrar_note"))
    if reg_note:
        _spacer(doc)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(reg_note)
        _apply_font(run, size_pt=BODY_SIZE)

    # ── 16. DISCLAIMER ────────────────────────────────────────────────────────
    disclaimer = _val(data.get("disclaimer"))
    if disclaimer:
        _spacer(doc)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(f"⚠  {disclaimer}")
        _apply_font(run, size_pt=BODY_SIZE)

    doc.save(str(output_path))


# ══════════════════════════════════════════════════════════════════════════════
#  PLOT DEED BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def _build_plot_docx(data: dict, output_path: Path):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.5)
        sec.right_margin  = Cm(2.5)

    hdr  = data.get("header",        {})
    v    = data.get("vendor",        {})
    pur  = data.get("purchaser",     {})
    con  = data.get("consideration", {})
    prop = data.get("property",      {})
    wit  = data.get("witnesses",     [{}, {}])

    # ── 1. TITLE ──────────────────────────────────────────────────────────────
    _title_para(doc, data.get("title", "சுத்த விக்கிரையப் பத்திரம்"))
    _title_para(doc, data.get("subtitle", "ABSOLUTE SALE DEED — PLOT / VACANT LAND"))
    _spacer(doc)

    # ── 2. OPENING PARAGRAPH ─────────────────────────────────────────────────
    date_day   = _req(hdr.get("date_day"))
    date_month = _req(hdr.get("date_month"))
    date_year  = _req(hdr.get("date_year"))

    date_str = f"{date_day}ம் {date_month} மாதம் {date_year}ம் தேதி,"

    pur_text = _party_text(pur, deed_type="plot")
    v_text   = _party_text(v,   deed_type="plot")

    opening = (
        f"{date_str}\n\n"
        f"{pur_text} அவர்களுக்கு,\n\n"
        f"{v_text} ஆகிய நான் அடியிற்கண்ட சாட்சிகள் முன்னிலையில் "
        f"மனப்பூர்வமாய் சம்மதித்து எழுதிக் கொள்ளும் காலிமனை "
        f"சுத்த விக்கிரையப் பத்திரம் என்னவென்றால்,"
    )
    _body_para(doc, opening, space_before=8, space_after=6, first_line_indent=False)

    # ── 3. ALL CLAUSES ───────────────────────────────────────────────────────
    for key in [
        "ownership_clause",
        "sale_clause",
        "possession_clause",
        "encumbrance_clause",
        "tax_clause",
        "patta_clause",
        "document_clause",
        "relinquish_clause",
    ]:
        val = _val(data.get(key))
        if val:
            _body_para(doc, val)

    # ── 4. CLOSING TEXT ───────────────────────────────────────────────────────
    _body_para(doc, _val(data.get("closing_text")), space_before=8, space_after=8)

    # ── 5. PROPERTY SCHEDULE ──────────────────────────────────────────────────
    _body_para(doc, "சொத்து விவரம்", bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=6, space_after=2, first_line_indent=False)

    prop_parts = []
    for label, key in [
        ("கதவு எண்",          "door_no"),      # optional
        ("வார்டு எண்",         "ward_no"),      # optional
        ("தொகுதி / Plot எண்", "plot_no"),      # optional
        ("தெரு",               "street"),       # optional
        ("பகுதி",              "area"),
        ("தாலுக்கா",           "taluk"),        # optional
        ("மாவட்டம்",           "district"),
    ]:
        val = _val(prop.get(key))
        if val:
            prop_parts.append(f"{label}: {val}")

    extent = _val(prop.get("extent_sqft"))
    if extent:
        prop_parts.append(f"மொத்த பரப்பு: {extent} Sq.ft")

    # Boundaries — optional individually
    boundary_parts = []
    for label, key in [
        ("கிழக்கு", "boundary_east"),
        ("மேற்கு",  "boundary_west"),
        ("வடக்கு",  "boundary_north"),
        ("தெற்கு",  "boundary_south"),
    ]:
        val = _val(prop.get(key))
        if val:
            boundary_parts.append(f"{label}: {val}")
    if boundary_parts:
        prop_parts.append("எல்லைகள் — " + ", ".join(boundary_parts))

    if prop_parts:
        _body_para(doc, ", ".join(prop_parts) + ".",
                   space_before=2, space_after=6, first_line_indent=False)

    # ── 6. WITNESSES ─────────────────────────────────────────────────────────
    _spacer(doc)
    w1 = wit[0] if len(wit) > 0 else {}
    w2 = wit[1] if len(wit) > 1 else {}

    for idx, w in enumerate([w1, w2], start=1):
        label   = w.get("label", f"சாட்சி {idx}")
        name    = _req(w.get("name"))
        address = _req(w.get("address"))
        aadhaar = _val(w.get("aadhaar"))          # optional
        line    = f"{label} : {name}, {address}"
        if aadhaar:
            line += f", ஆதார்: {aadhaar}"
        line += ",   கையொப்பம் : _______________"
        _sig_para(doc, line)

    _spacer(doc)

    # ── 7. SIGNATURES ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f"விற்பனையாளர் : {_req(v.get('name'))}")
    _apply_font(run, size_pt=BODY_SIZE)
    tab = p.add_run("\t\t\t\t\t")
    _apply_font(tab, size_pt=BODY_SIZE)
    run2 = p.add_run(f"கொள்முதலாளர் : {_req(pur.get('name'))}")
    _apply_font(run2, size_pt=BODY_SIZE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.first_line_indent = Cm(0)
    s1 = p2.add_run("கையொப்பம் / கட்டை விரல் ரேகை : _______________")
    _apply_font(s1, size_pt=BODY_SIZE)
    tab2 = p2.add_run("\t\t\t")
    _apply_font(tab2, size_pt=BODY_SIZE)
    s2 = p2.add_run("கையொப்பம் / கட்டை விரல் ரேகை : _______________")
    _apply_font(s2, size_pt=BODY_SIZE)

    # ── 8. LEGAL NOTES ────────────────────────────────────────────────────────
    legal_notes = _val(data.get("legal_notes"))
    if legal_notes:
        _spacer(doc)
        _body_para(doc, f"சட்ட குறிப்புகள் (Legal Notes) : {legal_notes}",
                   space_before=10, space_after=4, first_line_indent=False)

    doc.save(str(output_path))


# ══════════════════════════════════════════════════════════════════════════════
#  HANDLER
# ══════════════════════════════════════════════════════════════════════════════

def _count_unfilled(skeleton: dict) -> int:
    count = 0
    def _walk(obj):
        nonlocal count
        if isinstance(obj, str):
            if obj.startswith("{{") and obj.endswith("}}"):
                count += 1
        elif isinstance(obj, dict):
            for v in obj.values():
                _walk(v)
        elif isinstance(obj, list):
            for item in obj:
                _walk(item)
    _walk(skeleton)
    return count


async def handle(arguments: dict) -> list[TextContent]:
    filled_skeleton = arguments.get("filled_skeleton", {})
    prefix          = arguments.get("filename_prefix", "deed")
    deed_type       = filled_skeleton.get("type", "plot")

    unfilled_count = _count_unfilled(filled_skeleton)
    if unfilled_count > 10:
        return [TextContent(
            type="text",
            text=json.dumps({
                "success":   False,
                "filename":  None,
                "file":      None,
                "error": (
                    f"Skeleton has {unfilled_count} unfilled {{{{PLACEHOLDER}}}} values. "
                    "fill_skeleton (CALL 6) was not called or its 'clean_skeleton' output "
                    "was not passed here."
                ),
                "message": (
                    f"தோல்வி: skeleton-ல் {unfilled_count} field-கள் fill ஆகவில்லை. "
                    "fill_skeleton -> clean_skeleton -> generate_docx sequence பின்பற்றவும்."
                ),
                "next_tool": None,
            }, ensure_ascii=False)
        )]

    safe_prefix = "".join(c if c.isalnum() and ord(c) < 128 else "_" for c in prefix)
    first_name  = (safe_prefix.strip("_").split("_")[0] or "deed").lower()
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    filename    = f"{first_name}_{timestamp}.docx"
    output_path = OUTPUT_DIR / filename

    try:
        if deed_type == "agriculture":
            _build_agriculture_docx(filled_skeleton, output_path)
        else:
            _build_plot_docx(filled_skeleton, output_path)

        _mem_store(filename, output_path.read_bytes())

        return [TextContent(
            type="text",
            text=json.dumps({
                "success":   True,
                "filename":  filename,
                "file":      str(output_path),
                "error":     None,
                "message":   f"✅ பத்திரம் தயாரிக்கப்பட்டது: {filename}\nகோப்பை பார்க்க list_output_files tool call செய்.",
                "next_tool": "list_output_files"
            }, ensure_ascii=False, indent=2)
        )]

    except Exception as e:
        return [TextContent(
            type="text",
            text=json.dumps({
                "success":  False,
                "filename": None,
                "file":     None,
                "error":    str(e),
                "message":  f"❌ DOCX generation failed: {e}",
                "next_tool": None
            }, ensure_ascii=False)
        )]
